"""
Document Parser - Converts PDFs and DOCX files into hierarchical tree indexes.
No chunking. Preserves full document structure.

OCR Pipeline (for scanned/legacy documents):
  Poppler (pdftoppm) → renders PDF pages to images
  Tesseract           → OCR reads the images into text

Auto-detection: if a PDF page yields fewer than MIN_CHARS_PER_PAGE characters
of native text, it is treated as scanned and run through OCR automatically.
Mixed documents (some native, some scanned) are handled page-by-page.
"""

import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional
import hashlib

# A page with fewer than this many characters is considered scanned/image-only
MIN_CHARS_PER_PAGE = 50


# ── Dependency checks ────────────────────────────────────────────────────────

def _check_poppler() -> bool:
    """Return True if pdftoppm (Poppler) is available on PATH."""
    return shutil.which("pdftoppm") is not None


def _check_tesseract() -> bool:
    """Return True if tesseract is available on PATH."""
    return shutil.which("tesseract") is not None


def ocr_status() -> dict:
    """Report availability of OCR dependencies."""
    return {
        "poppler": _check_poppler(),
        "tesseract": _check_tesseract(),
        "ocr_available": _check_poppler() and _check_tesseract(),
    }


def parse_document(file_path: str) -> dict:
    """Parse any supported document into a hierarchical tree index."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _make_node(title: str, content: str, level: int, page: Optional[int] = None) -> dict:
    node_id = hashlib.md5(f"{title}{content[:50]}".encode()).hexdigest()[:8]
    return {
        "id": node_id,
        "title": title,
        "content": content.strip(),
        "level": level,
        "page": page,
        "children": []
    }


def _parse_pdf(file_path: str) -> dict:
    """
    Parse PDF into hierarchical structure.

    Strategy:
    1. Try native text extraction with PyMuPDF.
    2. For each page with suspiciously little text (scanned/image),
       fall back to Poppler → Tesseract OCR.
    3. Build tree from TOC if available, else heuristic heading detection.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    filename = Path(file_path).stem
    total_pages = len(doc)

    has_poppler = _check_poppler()
    has_tesseract = _check_tesseract()
    ocr_available = has_poppler and has_tesseract

    pages_text = []
    ocr_pages = 0

    for page_num, page in enumerate(doc, 1):
        native_text = page.get_text("text").strip()

        if len(native_text) >= MIN_CHARS_PER_PAGE:
            # Native text is fine
            pages_text.append({"page": page_num, "text": native_text, "ocr": False})
        elif ocr_available:
            # Scanned page — use Poppler + Tesseract
            ocr_text = _ocr_page(file_path, page_num)
            pages_text.append({"page": page_num, "text": ocr_text, "ocr": True})
            ocr_pages += 1
        else:
            # No OCR available, keep whatever little text we have
            pages_text.append({"page": page_num, "text": native_text, "ocr": False})

    toc = doc.get_toc()
    doc.close()

    root = {
        "title": filename,
        "source": file_path,
        "type": "pdf",
        "total_pages": total_pages,
        "ocr_pages": ocr_pages,
        "scanned": ocr_pages > (total_pages * 0.5),  # majority scanned?
        "children": []
    }

    if toc and len(toc) > 2:
        root["children"] = _build_tree_from_toc(toc, pages_text)
    else:
        root["children"] = _build_tree_from_text(pages_text)

    return root


def _ocr_page(file_path: str, page_num: int, dpi: int = 300, lang: str = "eng") -> str:
    """
    OCR a single PDF page using Poppler + Tesseract.

    Steps:
      1. pdftoppm renders the page to a PNG image at `dpi` resolution
      2. tesseract reads the image and outputs plain text
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        img_prefix = str(tmp / "page")

        # Step 1: Poppler — render just this one page to PNG
        # pdftoppm -png -r 300 -f 5 -l 5 input.pdf /tmp/page
        result = subprocess.run(
            [
                "pdftoppm",
                "-png",
                "-r", str(dpi),
                "-f", str(page_num),
                "-l", str(page_num),
                file_path,
                img_prefix,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            return f"[OCR failed for page {page_num}: {result.stderr.strip()}]"

        # Find the rendered image (pdftoppm names it like page-0005.png)
        images = sorted(tmp.glob("*.png"))
        if not images:
            return f"[No image rendered for page {page_num}]"

        img_path = images[0]

        # Step 2: Tesseract — OCR the image
        # tesseract img.png stdout -l eng
        result = subprocess.run(
            ["tesseract", str(img_path), "stdout", "-l", lang, "--psm", "3"],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            return f"[Tesseract error on page {page_num}: {result.stderr.strip()}]"

        return result.stdout.strip()


def _ocr_full_pdf(file_path: str, dpi: int = 300, lang: str = "eng") -> list:
    """
    OCR all pages of a PDF at once (faster for fully-scanned documents).
    Returns list of {page, text} dicts.
    """
    pages_text = []

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        img_prefix = str(tmp / "page")

        # Render all pages
        subprocess.run(
            ["pdftoppm", "-png", "-r", str(dpi), file_path, img_prefix],
            capture_output=True, timeout=300, check=True,
        )

        images = sorted(tmp.glob("*.png"))

        for i, img_path in enumerate(images, 1):
            result = subprocess.run(
                ["tesseract", str(img_path), "stdout", "-l", lang, "--psm", "3"],
                capture_output=True, text=True, timeout=120,
            )
            text = result.stdout.strip() if result.returncode == 0 else ""
            pages_text.append({"page": i, "text": text, "ocr": True})

    return pages_text


def _build_tree_from_toc(toc: list, pages_text: list) -> list:
    """Build a tree from PDF table of contents."""
    nodes = []
    stack = []  # [(level, node)]

    for i, (level, title, page) in enumerate(toc):
        # Get content between this and next TOC entry
        next_page = toc[i + 1][2] if i + 1 < len(toc) else len(pages_text) + 1
        content = _extract_pages_content(pages_text, page, next_page)

        node = _make_node(title, content, level, page)

        if level == 1:
            nodes.append(node)
            stack = [(1, node)]
        else:
            # Find parent
            while stack and stack[-1][0] >= level:
                stack.pop()
            if stack:
                stack[-1][1]["children"].append(node)
            else:
                nodes.append(node)
            stack.append((level, node))

    return nodes


def _extract_pages_content(pages_text: list, start_page: int, end_page: int) -> str:
    """Extract text from a range of pages."""
    content_parts = []
    for p in pages_text:
        if start_page <= p["page"] < end_page:
            content_parts.append(p["text"])
    return "\n".join(content_parts)[:3000]  # cap per section


def _build_tree_from_text(pages_text: list) -> list:
    """Heuristic heading detection when no TOC available."""
    heading_patterns = [
        (1, re.compile(r'^(CHAPTER\s+\d+|SECTION\s+\d+|PART\s+[IVX]+)\b', re.IGNORECASE | re.MULTILINE)),
        (1, re.compile(r'^\d+\.\s+[A-Z][A-Za-z\s]{3,60}$', re.MULTILINE)),
        (2, re.compile(r'^\d+\.\d+\s+[A-Z][A-Za-z\s]{3,60}$', re.MULTILINE)),
        (3, re.compile(r'^\d+\.\d+\.\d+\s+[A-Za-z\s]{3,60}$', re.MULTILINE)),
    ]

    all_sections = []
    for page_data in pages_text:
        text = page_data["text"]
        page = page_data["page"]
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue
            for level, pattern in heading_patterns:
                if pattern.match(line):
                    all_sections.append({
                        "title": line,
                        "level": level,
                        "page": page,
                        "content": ""
                    })
                    break

    # If very few headings found, just split by page groups
    if len(all_sections) < 3:
        return _build_page_groups(pages_text)

    # Assign content to each section
    nodes = []
    for i, sec in enumerate(all_sections):
        next_page = all_sections[i + 1]["page"] if i + 1 < len(all_sections) else 999
        content = _extract_pages_content(pages_text, sec["page"], next_page)
        node = _make_node(sec["title"], content, sec["level"], sec["page"])

        if sec["level"] == 1:
            nodes.append(node)
        elif nodes:
            nodes[-1]["children"].append(node)

    return nodes


def _build_page_groups(pages_text: list, group_size: int = 5) -> list:
    """Fallback: group pages into sections."""
    nodes = []
    for i in range(0, len(pages_text), group_size):
        group = pages_text[i:i + group_size]
        start_page = group[0]["page"]
        end_page = group[-1]["page"]
        content = "\n".join(p["text"] for p in group)
        node = _make_node(
            f"Pages {start_page}–{end_page}",
            content,
            1,
            start_page
        )
        nodes.append(node)
    return nodes


def _parse_docx(file_path: str) -> dict:
    """Parse DOCX into hierarchical structure using python-docx."""
    from docx import Document

    doc = Document(file_path)
    filename = Path(file_path).stem

    root = {
        "title": filename,
        "source": file_path,
        "type": "docx",
        "children": []
    }

    current_h1 = None
    current_h2 = None
    current_content = []

    def flush_content(node):
        if node and current_content:
            node["content"] = (node.get("content", "") + "\n" + "\n".join(current_content)).strip()

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if "Heading 1" in style:
            flush_content(current_h2 or current_h1)
            current_content = []
            current_h1 = _make_node(text, "", 1)
            current_h2 = None
            root["children"].append(current_h1)

        elif "Heading 2" in style:
            flush_content(current_h2)
            current_content = []
            current_h2 = _make_node(text, "", 2)
            if current_h1:
                current_h1["children"].append(current_h2)
            else:
                root["children"].append(current_h2)

        elif "Heading 3" in style:
            flush_content(current_h2)
            current_content = []
            h3 = _make_node(text, "", 3)
            target = current_h2 or current_h1
            if target:
                target["children"].append(h3)
            else:
                root["children"].append(h3)

        else:
            current_content.append(text)

    flush_content(current_h2 or current_h1)

    # If no headings found, just dump all text into root
    if not root["children"]:
        root["content"] = "\n".join(
            p.text for p in doc.paragraphs if p.text.strip()
        )[:5000]

    return root


def _parse_txt(file_path: str) -> dict:
    """Parse plain text file."""
    filename = Path(file_path).stem
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    return {
        "title": filename,
        "source": file_path,
        "type": "txt",
        "content": text[:10000],
        "children": []
    }


def _get_children(node: dict) -> list:
    return node.get("children") or node.get("nodes") or []

def _get_node_id(node: dict) -> Optional[str]:
    return node.get("id") or node.get("node_id")

def _get_page(node: dict) -> Optional[int]:
    # Explicit None check: avoids falsy bug where page=0 falls through to start_index
    p = node.get("page")
    if p is not None:
        return p
    return node.get("start_index")

def tree_to_outline(node: dict, depth: int = 0) -> str:
    indent = "  " * depth
    node_id = _get_node_id(node)
    page = _get_page(node)
    page_info = f" [p.{page}]" if page else ""
    id_info = f" [{node_id}]" if node_id else ""
    summary = node.get("summary", "")
    summary_hint = f" — {summary[:80]}" if summary else ""
    level = node.get("level", depth + 1)
    line = f"{indent}{'#' * max(1, level)} {node.get('title', 'Untitled')}{page_info}{id_info}{summary_hint}"
    lines = [line]
    for child in _get_children(node):
        lines.append(tree_to_outline(child, depth + 1))
    return "\n".join(lines)

def get_node_by_id(tree: dict, node_id: str) -> Optional[dict]:
    if _get_node_id(tree) == node_id:
        return tree
    for child in _get_children(tree):
        result = get_node_by_id(child, node_id)
        if result:
            return result
    return None


def save_index(tree: dict, output_path: str):
    """Save the tree index to a JSON file."""
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(tree, f, indent=2, ensure_ascii=False)


def load_index(index_path: str) -> dict:
    """Load a tree index from a JSON file."""
    with open(index_path, "r", encoding="utf-8") as f:
        return json.load(f)